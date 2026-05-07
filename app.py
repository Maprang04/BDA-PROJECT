import os
import streamlit as st
from docx import Document
from pypdf import PdfReader

from langchain.schema import Document as LangchainDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS


# =========================
# PAGE CONFIG
# =========================

st.set_page_config(
    page_title="MFU Grant Assistant",
    page_icon="📚",
    layout="wide"
)


# =========================
# GROUP INFORMATION
# =========================

GROUP_NO = "BDA_Project2_GroupNo1"

GROUP_MEMBERS = [
    {"student_id": "6631501120", "Achira Lueablae": "Student Name 1"},
    {"student_id": "65xxxxxxxx", "name": "Student Name 2"},
    {"student_id": "65xxxxxxxx", "name": "Student Name 3"},
]


# =========================
# DATASET PATH
# =========================

DATASET_DIR = "/content/sample_data"


# =========================
# READ DOCX
# =========================

def read_docx(file_path):
    doc = Document(file_path)
    text = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text.strip())

    return "\n".join(text)


# =========================
# READ PDF
# =========================

def read_pdf(file_path):
    text = []

    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    except Exception as e:
        text.append(f"Cannot read PDF: {e}")

    return "\n".join(text)


# =========================
# IMAGE KNOWLEDGE
# เพราะรูปภาพ Streamlit/RAG อ่านข้อความเองไม่ได้ดี
# จึงสรุปข้อความจากภาพเป็น text knowledge
# =========================

IMAGE_KNOWLEDGE = """
ขั้นตอนทุนประเภทที่ 1: ทุนสนับสนุนการจัดทำตำรา หนังสือ และงานแปลและเรียบเรียง เพื่อขอตำแหน่งทางวิชาการ
1. ผู้เขียนส่งผลงานและแบบคำขอรับทุนผ่านที่ประชุมสำนักวิชา
2. สำนักวิชาให้ความเห็นชอบและเสนอรายชื่อ
3. คณะกรรมการพิจารณาผลงานเบื้องต้น และคัดเลือกผู้ทรงคุณวุฒิ 2 คน
4. ผู้ทรงคุณวุฒิประเมินคุณภาพผลงาน
5. ถ้าผ่านระดับดีขึ้นไป ส่งแก้ไขและรออนุมัติทุน
6. ถ้าไม่ผ่านหรือต่ำกว่าระดับดี ต้องส่งแก้ไข และอาจขอพิจารณาใหม่ในปีงบประมาณหน้า
7. คณะกรรมการอนุมัติจัดสรรทุน
8. MLii ออกแบบรูปเล่ม ลงนามสัญญา และนำขึ้นจำหน่าย
9. ผู้เขียนได้รับค่าลิขสิทธิ์ 30% ของราคาจริงคูณด้วยจำนวนที่จำหน่ายได้

ขั้นตอนทุนประเภทที่ 2: ทุนจัดทำตำรา หนังสือ ประเภท eBook เพื่อการจำหน่าย
1. MLii ประชาสัมพันธ์โครงการไปยังสำนักวิชาและหน่วยงาน
2. ผู้เขียนตรวจสอบผลงานและเสนอคำขอรับการสนับสนุน พร้อมเอกสารประกอบตามแบบฟอร์ม
3. MLii รวบรวมและตรวจสอบความถูกต้องของเอกสาร
4. คณะกรรมการพิจารณาผลงานและคัดเลือกผู้ทรงคุณวุฒิ
5. ผู้เขียนประสานงานผู้ทรงคุณวุฒิเพื่อประเมินผลงาน
6. ถ้าผลประเมินต่ำกว่าระดับดี ผู้เขียนต้องปรับแก้ตามข้อเสนอแนะและยื่นขอรับการสนับสนุนใหม่อีกครั้ง
7. ถ้าผลประเมินระดับดีขึ้นไป ถือว่าผ่าน
8. MLii เสนอผลให้คณะกรรมการอนุมัติ
9. MLii นำผลงานจัดทำรูปเล่ม eBook และจัดทำสัญญาเพื่อการจำหน่าย
10. ผู้จำหน่ายนำผลงาน eBook ขึ้นจำหน่ายบนแพลตฟอร์ม
11. MLii เบิกจ่ายรายได้จากการจำหน่ายให้ผู้ขายจำนวนร้อยละ 80 หลังจากหักค่าใช้จ่ายทั้งหมด
"""


# =========================
# LOAD DOCUMENTS
# =========================

def load_documents():
    documents = []

    if not os.path.exists(DATASET_DIR):
        return documents

    for root, dirs, files in os.walk(DATASET_DIR):
        for file in files:
            file_path = os.path.join(root, file)
            lower_file = file.lower()

            if lower_file.endswith(".docx"):
                content = read_docx(file_path)
                documents.append(
                    LangchainDocument(
                        page_content=content,
                        metadata={"source": file}
                    )
                )

            elif lower_file.endswith(".pdf"):
                content = read_pdf(file_path)
                documents.append(
                    LangchainDocument(
                        page_content=content,
                        metadata={"source": file}
                    )
                )

    documents.append(
        LangchainDocument(
            page_content=IMAGE_KNOWLEDGE,
            metadata={"source": "สรุปจากภาพขั้นตอนทุนประเภทที่ 1 และ 2"}
        )
    )

    return documents


# =========================
# CREATE VECTORSTORE
# =========================

@st.cache_resource
def create_vectorstore():
    documents = load_documents()

    if len(documents) == 0:
        return None

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=150
    )

    chunks = text_splitter.split_documents(documents)

    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
    )

    vectorstore = FAISS.from_documents(chunks, embeddings)

    return vectorstore


# =========================
# RETRIEVE ANSWER
# =========================

def retrieve_answer(question, vectorstore):
    docs = vectorstore.similarity_search(question, k=4)

    if not docs:
        return "ไม่พบข้อมูลในเอกสารที่ให้", []

    answer = ""
    for i, doc in enumerate(docs, start=1):
        answer += f"\n\nข้อมูลอ้างอิงที่ {i}:\n"
        answer += doc.page_content[:1200]
        answer += f"\n\nแหล่งที่มา: {doc.metadata.get('source', 'Unknown')}\n"

    return answer, docs


# =========================
# UI
# =========================

st.title("📚 MFU Grant Assistant using RAG")
st.write("AI Chatbot สำหรับตอบคำถามเกี่ยวกับทุนตำรา หนังสือ eBook และการขอรับทุนสนับสนุน")

st.divider()

st.subheader("Project Information")
st.markdown(f"**Group No.: {GROUP_NO}**")

st.write("Group Members")

for member in GROUP_MEMBERS:
    st.write(f"- {member['student_id']} : {member['name']}")

st.divider()

st.subheader("About this AI Project")

st.markdown("""
ระบบนี้พัฒนาโดยใช้แนวคิด Retrieval-Augmented Generation หรือ RAG 
เพื่อค้นคืนข้อมูลจากเอกสาร Dataset ที่เกี่ยวข้องกับทุนตำรา หนังสือ และ eBook 
จากนั้นนำข้อมูลที่เกี่ยวข้องมาตอบคำถามของผู้ใช้
""")

st.divider()

st.subheader("Ask Questions")

question = st.text_input(
    "พิมพ์คำถามเกี่ยวกับทุนตำรา เช่น ทุนประเภทที่ 1 กับ 2 ต่างกันอย่างไร?"
)

vectorstore = create_vectorstore()

if question:
    if vectorstore is None:
        st.error("ไม่พบไฟล์ในโฟลเดอร์ dataset กรุณาตรวจสอบไฟล์อีกครั้ง")
    else:
        with st.spinner("กำลังค้นหาคำตอบจาก Dataset..."):
            answer, source_docs = retrieve_answer(question, vectorstore)

        st.subheader("Answer")
        st.write(answer)

        st.subheader("Retrieved Sources")
        for doc in source_docs:
            st.info(f"Source: {doc.metadata.get('source', 'Unknown')}")
